"""
docx_export.py – Word-Export im LRA Rosenheim Corporate Design
"""
import io
import os
import re
from datetime import date
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "lra_logo.jpg")
ORANGE    = RGBColor(0xFF, 0x66, 0x00)
GREY      = RGBColor(0x80, 0x80, 0x80)
BLACK     = RGBColor(0x1E, 0x1E, 0x1E)

DOC_TYPE_LABELS = {
    "aktenvermerk":          "AKTENVERMERK",
    "telefonnotiz":          "TELEFONNOTIZ",
    "besprechungsprotokoll": "BESPRECHUNGSPROTOKOLL",
}
DOC_FILENAMES = {
    "aktenvermerk":          "Aktenvermerk",
    "telefonnotiz":          "Telefonnotiz",
    "besprechungsprotokoll": "Besprechungsprotokoll",
}


# ---------------------------------------------------------------------------
# Hilfsfunktionen
# ---------------------------------------------------------------------------

def _run(para, text, size=11, bold=False, color=None, name="Arial"):
    r = para.add_run(text)
    r.font.name = name
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    return r


def _invisible_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _orange_line_para(container):
    p   = container.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:color"), "FF6600")
    bot.set(qn("w:space"), "4")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p


def _hint_box(doc, text):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "FFF4E6")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:color"), "FF6600")
    left.set(qn("w:space"), "4")
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent  = Mm(4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    _run(p, text, size=9, color=RGBColor(0x99, 0x44, 0x00))
    return p


def _add_footer(doc, status_text):
    sec    = doc.sections[0]
    footer = sec.footer
    p      = footer.paragraphs[0]
    p.clear()

    text_w_twips = int((sec.page_width - sec.left_margin - sec.right_margin) / 635)

    pPr = p._p.get_or_add_pPr()
    # Obere Trennlinie grau
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "4")
    top.set(qn("w:color"), "808080")
    top.set(qn("w:space"), "4")
    pBdr.append(top)
    pPr.append(pBdr)
    # Rechts-Tabstopp
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(text_w_twips))
    tabs.append(tab)
    pPr.append(tabs)

    _run(p, status_text, size=8, color=GREY)
    p.add_run("\t")
    _run(p, "Seite ", size=8, color=GREY)

    # Seitenzahl-Feld
    r_begin = p.add_run()
    r_begin.font.name = "Arial"; r_begin.font.size = Pt(8); r_begin.font.color.rgb = GREY
    fc_b = OxmlElement("w:fldChar"); fc_b.set(qn("w:fldCharType"), "begin")
    r_begin._r.append(fc_b)

    r_instr = p.add_run()
    r_instr.font.name = "Arial"; r_instr.font.size = Pt(8); r_instr.font.color.rgb = GREY
    ins = OxmlElement("w:instrText"); ins.text = " PAGE "
    r_instr._r.append(ins)

    r_end = p.add_run()
    r_end.font.name = "Arial"; r_end.font.size = Pt(8); r_end.font.color.rgb = GREY
    fc_e = OxmlElement("w:fldChar"); fc_e.set(qn("w:fldCharType"), "end")
    r_end._r.append(fc_e)


def _smart_quotes(text: str) -> str:
    text = re.sub(r'"([^"]*)"', "„\\1“", text)
    text = re.sub(r"'([^']*)'", "‚\\1‘", text)
    return text


def _parse_text(text: str):
    """Trennt Metadaten-Zeilen (Key: Value) und Fließtext."""
    title_skipped = False
    meta_pairs    = []
    body_lines    = []
    in_meta       = True

    for line in text.split("\n"):
        s = line.strip()
        if not s:
            if not in_meta:
                body_lines.append("")
            continue
        if len(s) > 2 and all(c in "=-" for c in s):
            continue
        if not title_skipped:
            if s.upper() == s and len(s) >= 3:
                title_skipped = True
            continue
        if in_meta and ":" in s and not s[0].isdigit():
            colon = s.index(":")
            key   = s[:colon].strip()
            val   = s[colon + 1:].strip()
            clean = key.replace(" ", "").replace("/", "").replace("-", "")
            if len(key) < 30 and clean.isalpha():
                meta_pairs.append((key, val))
                continue
        in_meta = False
        body_lines.append(s)

    return meta_pairs, body_lines


# ---------------------------------------------------------------------------
# Öffentliche Funktion
# ---------------------------------------------------------------------------

def create_docx(text: str, doc_type: str, metadata: dict) -> bytes:
    doc   = Document()
    datum = metadata.get("datum", date.today().strftime("%d.%m.%Y"))

    # Seiteneinrichtung A4, Ränder 25/20/25/25 mm
    sec = doc.sections[0]
    sec.page_width    = Mm(210)
    sec.page_height   = Mm(297)
    sec.top_margin    = Mm(25)
    sec.bottom_margin = Mm(25)
    sec.left_margin   = Mm(25)
    sec.right_margin  = Mm(20)

    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    # --- Logo-Header-Tabelle (Logo links | Gz./Datum rechts) ---
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment        = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Mm(80)
    tbl.columns[1].width = Mm(85)

    logo_cell = tbl.cell(0, 0)
    logo_para = logo_cell.paragraphs[0]
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after  = Pt(0)
    if os.path.exists(LOGO_PATH):
        logo_para.add_run().add_picture(LOGO_PATH, height=Mm(18))

    gz_cell = tbl.cell(0, 1)
    gz_para  = gz_cell.paragraphs[0]
    gz_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    gz_para.paragraph_format.space_before = Pt(0)
    gz_para.paragraph_format.space_after  = Pt(0)
    _run(gz_para, f"Gz.: ____\n{datum}", size=9, color=GREY)

    for cell in tbl.rows[0].cells:
        _invisible_cell_borders(cell)

    # --- Orangene Trennlinie ---
    _orange_line_para(doc)

    # --- Dokumenttitel 16 pt fett ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after  = Pt(4)
    _run(p_title, DOC_TYPE_LABELS.get(doc_type, "DOKUMENT"), size=16, bold=True)

    _orange_line_para(doc)

    # --- KI-Hinweis-Box ---
    _hint_box(doc, "KI-generierter Entwurf – vor Ablage fachlich prüfen")

    # --- Metadaten-Tabelle (2-spaltig, unsichtbar) ---
    meta_pairs, body_lines = _parse_text(text)

    if meta_pairs:
        mtbl = doc.add_table(rows=len(meta_pairs), cols=2)
        mtbl.alignment        = WD_TABLE_ALIGNMENT.LEFT
        mtbl.columns[0].width = Mm(40)
        mtbl.columns[1].width = Mm(125)
        for i, (key, val) in enumerate(meta_pairs):
            kc = mtbl.cell(i, 0)
            vc = mtbl.cell(i, 1)
            kp = kc.paragraphs[0]
            vp = vc.paragraphs[0]
            kp.paragraph_format.space_before = Pt(1)
            kp.paragraph_format.space_after  = Pt(1)
            vp.paragraph_format.space_before = Pt(1)
            vp.paragraph_format.space_after  = Pt(1)
            _run(kp, key + ":", size=10, bold=True, color=GREY)
            _run(vp, _smart_quotes(val), size=10, color=BLACK)
            for cell in (kc, vc):
                _invisible_cell_borders(cell)

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)

    # --- Fließtext mit Überschriften ---
    for line in body_lines:
        s = line.strip()
        if not s:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            continue

        # Nummerierte Abschnittsüberschrift (z. B. "1. ANLASS UND SACHVERHALT")
        if s[0].isdigit() and len(s) > 2 and s[1] in (".", " "):
            rest = s.split(".", 1)[-1].strip() if "." in s[:3] else s
            if rest.isupper():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after  = Pt(3)
                _run(p, s, size=12, bold=True, color=ORANGE)
                pPr  = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bot  = OxmlElement("w:bottom")
                bot.set(qn("w:val"),   "single")
                bot.set(qn("w:sz"),    "4")
                bot.set(qn("w:color"), "FF6600")
                bot.set(qn("w:space"), "2")
                pBdr.append(bot)
                pPr.append(pBdr)
                continue

        # Normaler Absatz
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        _run(p, _smart_quotes(s), size=11, color=BLACK)

    # --- Fußzeile ---
    _add_footer(
        doc,
        "Dokumentenstatus: Entwurf | KI-generiert – vor Ablage fachlich prüfen"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
